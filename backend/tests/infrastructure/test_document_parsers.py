from io import BytesIO

import pytest
from docx import Document
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

from backend.app.infrastructure.document_parsers import DocumentParseError, DocumentTextParser


def docx_bytes() -> bytes:
    document = Document()
    document.add_heading("支付需求", level=1)
    document.add_paragraph("系统必须支持退款。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "状态"
    table.cell(0, 1).text = "说明"
    table.cell(1, 0).text = "成功"
    table.cell(1, 1).text = "退款完成"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def pdf_bytes(*, text: str | None = None, encrypted: bool = False) -> bytes:
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    if text is not None:
        font = DictionaryObject(
            {
                NameObject("/Type"): NameObject("/Font"),
                NameObject("/Subtype"): NameObject("/Type1"),
                NameObject("/BaseFont"): NameObject("/Helvetica"),
            }
        )
        page[NameObject("/Resources")] = DictionaryObject(
            {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font})}
        )
        stream = DecodedStreamObject()
        stream.set_data(f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode("ascii"))
        page[NameObject("/Contents")] = writer._add_object(stream)
    if encrypted:
        writer.encrypt("secret")
    output = BytesIO()
    writer.write(output)
    return output.getvalue()


def test_docx_parser_preserves_paragraphs_and_tables() -> None:
    parsed = DocumentTextParser().parse(docx_bytes(), ".docx")

    assert "支付需求" in parsed
    assert "系统必须支持退款。" in parsed
    assert "状态\t说明" in parsed
    assert "成功\t退款完成" in parsed


def test_docx_parser_rejects_excessive_uncompressed_container() -> None:
    with pytest.raises(DocumentParseError, match="document_too_complex"):
        DocumentTextParser(max_docx_uncompressed_bytes=10).parse(docx_bytes(), ".docx")


def test_pdf_parser_adds_page_boundaries() -> None:
    parsed = DocumentTextParser().parse(pdf_bytes(text="Refund required"), ".pdf")

    assert parsed == "[第 1 页]\nRefund required"


def test_text_parser_builds_stable_line_chunks_and_offsets() -> None:
    result = DocumentTextParser().parse_document(
        "# 支付需求\n必须支持退款。\n\n## 异常流\n余额不足时拒绝退款。".encode(),
        ".md",
    )

    assert result.text == "# 支付需求\n必须支持退款。\n\n## 异常流\n余额不足时拒绝退款。"
    assert [chunk.source_type for chunk in result.chunks] == ["lines", "lines"]
    assert [(chunk.source_start, chunk.source_end) for chunk in result.chunks] == [
        (1, 2),
        (4, 5),
    ]
    assert result.text[result.chunks[1].start_offset : result.chunks[1].end_offset] == (
        "## 异常流\n余额不足时拒绝退款。"
    )


def test_binary_parser_builds_block_and_page_chunks() -> None:
    docx = DocumentTextParser().parse_document(docx_bytes(), ".docx")
    pdf = DocumentTextParser().parse_document(pdf_bytes(text="Refund required"), ".pdf")

    assert [chunk.source_type for chunk in docx.chunks] == ["block", "block", "block"]
    assert [(chunk.source_start, chunk.source_end) for chunk in docx.chunks] == [
        (1, 1),
        (2, 2),
        (3, 3),
    ]
    assert len(pdf.chunks) == 1
    assert pdf.chunks[0].source_type == "page"
    assert pdf.chunks[0].source_start == 1
    assert pdf.chunks[0].text == "[第 1 页]\nRefund required"


@pytest.mark.parametrize(
    ("content", "suffix", "reason"),
    [
        (b"not-a-docx", ".docx", "corrupt_document"),
        (b"not-a-pdf", ".pdf", "corrupt_document"),
        (pdf_bytes(), ".pdf", "empty_text"),
        (pdf_bytes(encrypted=True), ".pdf", "encrypted_pdf"),
    ],
)
def test_binary_parser_reports_safe_failure_reasons(
    content: bytes, suffix: str, reason: str
) -> None:
    with pytest.raises(DocumentParseError, match=reason):
        DocumentTextParser().parse(content, suffix)


def test_text_parser_rejects_damaged_utf8_and_empty_output() -> None:
    parser = DocumentTextParser()

    with pytest.raises(DocumentParseError, match="invalid_encoding"):
        parser.parse(b"\xff\xfe", ".txt")
    with pytest.raises(DocumentParseError, match="empty_text"):
        parser.parse(b"  \n", ".md")
    with pytest.raises(DocumentParseError, match="extracted_text_too_large"):
        DocumentTextParser(max_text_chars=5).parse(b"sixsix", ".txt")
